"""
Smart Document Generator v2
===========================
Uses improved NLP pipeline for better content extraction.
NOW WITH GENERATIVE CONTENT using offline AI models!

Generates structured documents with:
- Executive Summary
- Key Sentences  
- Key Concepts (with definitions)
- Study Questions (easy/medium/hard)
- Related Topics
- Action Items (for meetings)
- Complete Content

NEW - AI Enhanced Sections:
- Simplified Explanation
- Key Takeaways
- Real-world Examples
- FAQ (Auto-generated Q&A)
- Vocabulary with Definitions

Formats: Markdown, HTML, PDF, DOCX, TXT
"""
import re
from typing import Optional
from pathlib import Path
from datetime import datetime

# Import the improved analyzer
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from nlp.smart_analyzer import SmartAnalyzer, ContentAnalysis
from nlp.content_enhancer import ContentEnhancer, OfflineContentGenerator, GeneratedContent, get_content_enhancer


class SmartDocumentGenerator:
    """
    Generate structured documents from transcripts/text
    
    Uses improved NLP for better extraction of:
    - Key concepts with context
    - Study questions at different difficulty levels
    - Related topics for further study
    
    NEW: AI-powered content generation for:
    - Simplified explanations
    - Key takeaways
    - Real-world examples
    - Auto-generated FAQ
    - Vocabulary definitions
    
    Usage:
        generator = SmartDocumentGenerator()
        
        # Generate markdown (without AI enhancement)
        generator.generate("transcript.txt", "output.md")
        
        # Generate HTML with AI enhancement
        generator.generate("transcript.txt", "output.html", format="html", use_ai=True)
    """
    
    def __init__(self, use_ai: bool = False):
        """
        Initialize generator
        
        Args:
            use_ai: Whether to use AI for content generation
                    Requires: pip install transformers torch
        """
        self.analyzer = SmartAnalyzer()
        self.use_ai = use_ai
        self.enhancer = None
        
        if use_ai:
            self.enhancer = get_content_enhancer(use_ai=True)
    
    def generate(
        self,
        text: str,
        output_path: str,
        title: str = "EchoNotes Document",
        format: str = "auto",
        include_full_content: bool = True,
        use_ai: bool = None
    ) -> str:
        """
        Generate formatted document
        
        Args:
            text: Input text/transcript
            output_path: Output file path
            title: Document title
            format: Output format (auto, md, html, txt, pdf, docx)
            include_full_content: Whether to include full text
            use_ai: Override AI setting for this generation
            
        Returns:
            Path to generated file
        """
        # Analyze content
        print(f"🔍 Analyzing content...")
        analysis = self.analyzer.analyze(text, title)
        
        print(f"   ✓ Found {len(analysis.concepts)} key concepts")
        print(f"   ✓ Generated {len(analysis.questions)} study questions")
        print(f"   ✓ Identified {len(analysis.related_topics)} related topics")
        
        # AI Enhancement
        ai_content = None
        should_use_ai = use_ai if use_ai is not None else self.use_ai
        
        if should_use_ai:
            if self.enhancer is None:
                self.enhancer = get_content_enhancer(use_ai=True)
            ai_content = self.enhancer.enhance_content(text, title)
        
        # Determine format
        if format == "auto":
            ext = Path(output_path).suffix.lower()
            format = ext[1:] if ext else 'md'
        
        # Generate document
        if format == 'md':
            return self._generate_markdown(analysis, output_path, include_full_content, ai_content)
        elif format == 'html':
            return self._generate_html(analysis, output_path, include_full_content, ai_content)
        elif format == 'txt':
            return self._generate_text(analysis, output_path, include_full_content, ai_content)
        elif format == 'pdf':
            return self._generate_pdf(analysis, output_path, include_full_content, ai_content)
        elif format == 'docx':
            return self._generate_docx(analysis, output_path, include_full_content, ai_content)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def generate_from_file(
        self,
        input_path: str,
        output_path: Optional[str] = None,
        format: str = "md",
        use_ai: bool = False
    ) -> str:
        """Generate document from input file"""
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        if output_path is None:
            stem = Path(input_path).stem
            output_path = f"{stem}_notes.{format}"
        
        title = Path(input_path).stem.replace('_', ' ').title()
        return self.generate(text, output_path, title, format, use_ai=use_ai)
    
    def _generate_markdown(
        self,
        analysis: ContentAnalysis,
        path: str,
        include_full: bool,
        ai_content: Optional[GeneratedContent] = None
    ) -> str:
        """Generate Markdown document"""
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        
        lines = [
            f"# 📝 {analysis.title}",
            "",
            f"> **Generated:** {now}  ",
            f"> **Words:** {analysis.word_count} | **Reading Time:** {analysis.reading_time_minutes} min",
            "",
            "---",
            "",
            "## 📋 Executive Summary",
            "",
            analysis.executive_summary,
            "",
        ]
        
        # AI Enhanced: Simplified Explanation
        if ai_content and ai_content.simplified_explanation:
            lines.extend([
                "---",
                "",
                "## 🎯 Simple Explanation",
                "",
                ai_content.simplified_explanation,
                "",
            ])
        
        # Key Takeaways (AI or extracted)
        takeaways = ai_content.key_takeaways if ai_content else analysis.key_sentences[:5]
        lines.extend([
            "---",
            "",
            "## 🎯 Key Takeaways",
            "",
        ])
        for i, point in enumerate(takeaways, 1):
            lines.append(f"**{i}.** {point}")
        lines.append("")
        
        # Key Sentences
        lines.extend([
            "---",
            "",
            "## 🔑 Key Sentences",
            "",
        ])
        for i, sent in enumerate(analysis.key_sentences, 1):
            lines.append(f"{i}. {sent}")
        
        lines.extend([
            "",
            "---",
            "",
            "## 💡 Key Concepts",
            "",
        ])
        
        for concept in analysis.concepts:
            stars = int(concept.importance_score * 5)
            star_str = "⭐" * stars + "☆" * (5 - stars)
            lines.extend([
                f"### {concept.term}",
                f"**Importance:** {star_str}",
                "",
                f"> {concept.definition}",
                "",
            ])
        
        # AI Enhanced: Vocabulary
        if ai_content and ai_content.vocabulary:
            lines.extend([
                "---",
                "",
                "## 📚 Vocabulary",
                "",
                "| Term | Meaning |",
                "|------|---------|",
            ])
            for vocab in ai_content.vocabulary:
                term = vocab['term']
                meaning = vocab['meaning'][:100] + "..." if len(vocab['meaning']) > 100 else vocab['meaning']
                lines.append(f"| **{term}** | {meaning} |")
            lines.append("")
        
        # AI Enhanced: Examples
        if ai_content and ai_content.examples:
            lines.extend([
                "---",
                "",
                "## 💡 Real-World Examples",
                "",
            ])
            for i, example in enumerate(ai_content.examples, 1):
                lines.append(f"**Example {i}:** {example}")
                lines.append("")
        
        lines.extend([
            "---",
            "",
            "## ❓ Study Questions",
            "",
        ])
        
        difficulty_emoji = {'easy': '🟢', 'medium': '🟡', 'hard': '🔴'}
        
        for i, q in enumerate(analysis.questions, 1):
            emoji = difficulty_emoji.get(q.difficulty, '⚪')
            lines.extend([
                f"### {i}. {q.question}",
                f"- **Difficulty:** {emoji} {q.difficulty.title()}",
                f"- **Type:** {q.question_type.title()}",
            ])
            if q.answer_hint:
                lines.append(f"- **Hint:** *{q.answer_hint}*")
            lines.append("")
        
        # AI Enhanced: FAQ
        if ai_content and ai_content.faq:
            lines.extend([
                "---",
                "",
                "## ❔ Frequently Asked Questions",
                "",
            ])
            for i, qa in enumerate(ai_content.faq, 1):
                lines.extend([
                    f"**Q{i}: {qa['q']}**",
                    "",
                    f"> {qa['a']}",
                    "",
                ])
        
        lines.extend([
            "---",
            "",
            "## 🔗 Related Topics for Further Study",
            "",
        ])
        
        for topic in analysis.related_topics:
            lines.append(f"- 📚 {topic}")
        
        # Meeting-specific sections
        if analysis.action_items:
            lines.extend([
                "",
                "---",
                "",
                "## ✅ Action Items",
                "",
            ])
            for item in analysis.action_items:
                lines.append(f"- [ ] {item}")
        
        if analysis.decisions:
            lines.extend([
                "",
                "---",
                "",
                "## 🎯 Decisions Made",
                "",
            ])
            for decision in analysis.decisions:
                lines.append(f"- ✓ {decision}")
        
        if analysis.deadlines:
            lines.extend([
                "",
                "---",
                "",
                "## 📅 Deadlines",
                "",
            ])
            for deadline in analysis.deadlines:
                lines.append(f"- 🗓️ {deadline}")
        
        # Full content
        if include_full:
            lines.extend([
                "",
                "---",
                "",
                "## 📄 Complete Content",
                "",
                "<details>",
                "<summary>Click to expand full content</summary>",
                "",
                "```",
                analysis.executive_summary,  # Use clean version
                "```",
                "",
                "</details>",
            ])
        
        # AI indicator
        if ai_content:
            lines.extend([
                "",
                "---",
                "",
                "🤖 *This document includes AI-generated content for enhanced understanding*"
            ])
        
        lines.extend([
            "",
            "---",
            "",
            "*Generated by EchoNotes - Speech to Document System*"
        ])
        
        with open(path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        return path
    
    def _generate_html(
        self,
        analysis: ContentAnalysis,
        path: str,
        include_full: bool,
        ai_content: Optional[GeneratedContent] = None
    ) -> str:
        """Generate beautiful HTML document"""
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{analysis.title}</title>
    <style>
        * {{ box-sizing: border-box; margin: 0; padding: 0; }}
        body {{
            font-family: 'Segoe UI', system-ui, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }}
        .container {{
            max-width: 900px;
            margin: 0 auto;
            background: white;
            border-radius: 20px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            overflow: hidden;
        }}
        .header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 40px;
            text-align: center;
        }}
        .header h1 {{ font-size: 2.5em; margin-bottom: 10px; }}
        .header .meta {{ opacity: 0.9; font-size: 0.95em; }}
        .content {{ padding: 40px; }}
        .section {{
            margin-bottom: 40px;
            padding-bottom: 30px;
            border-bottom: 2px solid #eee;
        }}
        .section:last-child {{ border-bottom: none; }}
        .section-title {{
            display: flex;
            align-items: center;
            gap: 12px;
            color: #333;
            font-size: 1.5em;
            margin-bottom: 20px;
        }}
        .section-title .icon {{ font-size: 1.2em; }}
        .summary-box {{
            background: linear-gradient(135deg, #f5f7fa 0%, #e4e8eb 100%);
            padding: 25px;
            border-radius: 12px;
            font-size: 1.1em;
            line-height: 1.7;
            border-left: 5px solid #667eea;
        }}
        .key-sentence {{
            background: #f8f9fa;
            padding: 15px 20px;
            margin: 10px 0;
            border-radius: 8px;
            border-left: 4px solid #28a745;
            transition: transform 0.2s;
        }}
        .key-sentence:hover {{ transform: translateX(5px); }}
        .concept-card {{
            background: white;
            border: 2px solid #e9ecef;
            border-radius: 12px;
            padding: 20px;
            margin: 15px 0;
            transition: all 0.3s;
        }}
        .concept-card:hover {{
            border-color: #667eea;
            box-shadow: 0 5px 20px rgba(102, 126, 234, 0.2);
        }}
        .concept-title {{
            font-size: 1.2em;
            color: #333;
            margin-bottom: 8px;
        }}
        .concept-stars {{ color: #ffc107; letter-spacing: 2px; }}
        .concept-def {{
            color: #666;
            line-height: 1.6;
            margin-top: 10px;
            padding: 10px;
            background: #f8f9fa;
            border-radius: 6px;
        }}
        .question-card {{
            background: #fff3e0;
            border-radius: 12px;
            padding: 20px;
            margin: 15px 0;
        }}
        .question-text {{
            font-size: 1.1em;
            color: #333;
            margin-bottom: 12px;
        }}
        .question-meta {{
            display: flex;
            gap: 20px;
            font-size: 0.9em;
        }}
        .difficulty {{
            display: inline-block;
            padding: 4px 12px;
            border-radius: 20px;
            font-weight: 600;
        }}
        .difficulty.easy {{ background: #d4edda; color: #155724; }}
        .difficulty.medium {{ background: #fff3cd; color: #856404; }}
        .difficulty.hard {{ background: #f8d7da; color: #721c24; }}
        .hint {{
            color: #666;
            font-style: italic;
            margin-top: 10px;
            padding: 10px;
            background: rgba(255,255,255,0.7);
            border-radius: 6px;
        }}
        .topic-tags {{
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
        }}
        .topic-tag {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 10px 20px;
            border-radius: 25px;
            font-size: 0.95em;
            transition: transform 0.2s;
        }}
        .topic-tag:hover {{ transform: scale(1.05); }}
        .action-item {{
            display: flex;
            align-items: center;
            gap: 10px;
            padding: 12px 15px;
            background: #e8f5e9;
            border-radius: 8px;
            margin: 8px 0;
        }}
        .action-item input {{ width: 18px; height: 18px; }}
        .footer {{
            text-align: center;
            padding: 20px;
            background: #f8f9fa;
            color: #666;
            font-size: 0.9em;
        }}
        details {{
            background: #f8f9fa;
            border-radius: 8px;
            padding: 15px;
        }}
        summary {{
            cursor: pointer;
            font-weight: 600;
            color: #667eea;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>📝 {analysis.title}</h1>
            <div class="meta">
                <strong>Generated:</strong> {now} | 
                <strong>Words:</strong> {analysis.word_count} | 
                <strong>Reading Time:</strong> {analysis.reading_time_minutes} min
            </div>
        </div>
        
        <div class="content">
            <div class="section">
                <h2 class="section-title"><span class="icon">📋</span> Executive Summary</h2>
                <div class="summary-box">{analysis.executive_summary}</div>
            </div>
            
            <div class="section">
                <h2 class="section-title"><span class="icon">🔑</span> Key Sentences</h2>
"""
        
        for sent in analysis.key_sentences:
            html += f'                <div class="key-sentence">{sent}</div>\n'
        
        html += """
            </div>
            
            <div class="section">
                <h2 class="section-title"><span class="icon">💡</span> Key Concepts</h2>
"""
        
        for concept in analysis.concepts:
            stars = int(concept.importance_score * 5)
            star_str = "★" * stars + "☆" * (5 - stars)
            html += f"""                <div class="concept-card">
                    <div class="concept-title">{concept.term}</div>
                    <div class="concept-stars">{star_str}</div>
                    <div class="concept-def">{concept.definition}</div>
                </div>
"""
        
        html += """
            </div>
            
            <div class="section">
                <h2 class="section-title"><span class="icon">❓</span> Study Questions</h2>
"""
        
        for i, q in enumerate(analysis.questions, 1):
            html += f"""                <div class="question-card">
                    <div class="question-text"><strong>{i}.</strong> {q.question}</div>
                    <div class="question-meta">
                        <span class="difficulty {q.difficulty}">{q.difficulty.upper()}</span>
                        <span>Type: {q.question_type.title()}</span>
                    </div>
                    <div class="hint">💡 Hint: {q.answer_hint}</div>
                </div>
"""
        
        html += """
            </div>
            
            <div class="section">
                <h2 class="section-title"><span class="icon">🔗</span> Related Topics</h2>
                <div class="topic-tags">
"""
        
        for topic in analysis.related_topics:
            html += f'                    <span class="topic-tag">📚 {topic}</span>\n'
        
        html += """                </div>
            </div>
"""
        
        # Action items if present
        if analysis.action_items:
            html += """
            <div class="section">
                <h2 class="section-title"><span class="icon">✅</span> Action Items</h2>
"""
            for item in analysis.action_items:
                html += f"""                <div class="action-item">
                    <input type="checkbox">
                    <span>{item}</span>
                </div>
"""
            html += "            </div>\n"
        
        # ============ AI ENHANCED SECTIONS ============
        if ai_content:
            # Simple Explanation
            if ai_content.simplified_explanation:
                html += f"""
            <div class="section" style="background: linear-gradient(135deg, #e8f5e9 0%, #c8e6c9 100%); margin: 20px -40px; padding: 40px;">
                <h2 class="section-title"><span class="icon">🎯</span> Simple Explanation</h2>
                <div class="summary-box" style="background: white; border-left-color: #4caf50;">
                    {ai_content.simplified_explanation}
                </div>
            </div>
"""
            
            # Key Takeaways
            if ai_content.key_takeaways:
                html += """
            <div class="section">
                <h2 class="section-title"><span class="icon">📝</span> Key Takeaways</h2>
                <div style="display: grid; gap: 12px;">
"""
                for i, takeaway in enumerate(ai_content.key_takeaways, 1):
                    html += f"""                    <div style="background: #e3f2fd; padding: 15px 20px; border-radius: 8px; display: flex; gap: 12px; align-items: flex-start;">
                        <span style="background: #1976d2; color: white; width: 28px; height: 28px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-weight: bold; flex-shrink: 0;">{i}</span>
                        <span>{takeaway}</span>
                    </div>
"""
                html += """                </div>
            </div>
"""
            
            # Real-World Examples
            if ai_content.examples:
                html += """
            <div class="section">
                <h2 class="section-title"><span class="icon">💡</span> Real-World Examples</h2>
"""
                for i, example in enumerate(ai_content.examples, 1):
                    html += f"""                <div style="background: #f3e5f5; padding: 20px; border-radius: 12px; margin: 10px 0; border-left: 4px solid #9c27b0;">
                    <strong style="color: #7b1fa2;">Example {i}:</strong> {example}
                </div>
"""
                html += "            </div>\n"
            
            # FAQ
            if ai_content.faq:
                html += """
            <div class="section" style="background: #fff3e0; margin: 20px -40px; padding: 40px;">
                <h2 class="section-title"><span class="icon">❔</span> Frequently Asked Questions</h2>
"""
                for i, qa in enumerate(ai_content.faq, 1):
                    html += f"""                <div style="background: white; padding: 20px; border-radius: 12px; margin: 15px 0; box-shadow: 0 2px 8px rgba(0,0,0,0.1);">
                    <div style="font-weight: 600; color: #e65100; margin-bottom: 10px;">Q{i}: {qa['q']}</div>
                    <div style="color: #555; padding-left: 20px; border-left: 3px solid #ff9800;">{qa['a']}</div>
                </div>
"""
                html += "            </div>\n"
            
            # Vocabulary
            if ai_content.vocabulary:
                html += """
            <div class="section">
                <h2 class="section-title"><span class="icon">📚</span> Vocabulary</h2>
                <div style="display: grid; gap: 10px;">
"""
                for vocab in ai_content.vocabulary:
                    html += f"""                    <div style="background: #fafafa; padding: 15px; border-radius: 8px; border: 1px solid #e0e0e0;">
                        <strong style="color: #1976d2;">{vocab['term']}</strong>: {vocab['meaning']}
                    </div>
"""
                html += """                </div>
            </div>
"""
            
            # AI Badge
            html += """
            <div style="text-align: center; padding: 20px; margin-top: 20px;">
                <span style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 10px 25px; border-radius: 25px; font-size: 0.9em;">
                    🤖 AI-Enhanced Document
                </span>
            </div>
"""
        
        html += f"""
        </div>
        
        <div class="footer">
            Generated by EchoNotes - Speech to Document System
        </div>
    </div>
</body>
</html>"""
        
        with open(path, 'w', encoding='utf-8') as f:
            f.write(html)
        
        return path
    
    def _generate_text(
        self,
        analysis: ContentAnalysis,
        path: str,
        include_full: bool,
        ai_content: Optional[GeneratedContent] = None
    ) -> str:
        """Generate plain text document"""
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        
        lines = [
            "=" * 70,
            f"  {analysis.title}",
            "=" * 70,
            f"  Generated: {now}",
            f"  Words: {analysis.word_count} | Reading Time: {analysis.reading_time_minutes} min",
            "=" * 70,
            "",
            "EXECUTIVE SUMMARY",
            "-" * 50,
            analysis.executive_summary,
            "",
            "",
            "KEY SENTENCES",
            "-" * 50,
        ]
        
        for i, sent in enumerate(analysis.key_sentences, 1):
            lines.append(f"  {i}. {sent}")
        
        lines.extend([
            "",
            "",
            "KEY CONCEPTS",
            "-" * 50,
        ])
        
        for concept in analysis.concepts:
            stars = "*" * int(concept.importance_score * 5)
            lines.extend([
                f"  [{stars}] {concept.term}",
                f"      {concept.definition}",
                ""
            ])
        
        lines.extend([
            "",
            "STUDY QUESTIONS",
            "-" * 50,
        ])
        
        for i, q in enumerate(analysis.questions, 1):
            diff_marker = {'easy': '[E]', 'medium': '[M]', 'hard': '[H]'}.get(q.difficulty, '[?]')
            lines.extend([
                f"  {i}. {diff_marker} {q.question}",
                f"     Type: {q.question_type}",
                f"     Hint: {q.answer_hint}",
                ""
            ])
        
        lines.extend([
            "",
            "RELATED TOPICS",
            "-" * 50,
        ])
        
        for topic in analysis.related_topics:
            lines.append(f"  - {topic}")
        
        if analysis.action_items:
            lines.extend([
                "",
                "",
                "ACTION ITEMS",
                "-" * 50,
            ])
            for item in analysis.action_items:
                lines.append(f"  [ ] {item}")
        
        lines.extend([
            "",
            "=" * 70,
            "  Generated by EchoNotes",
            "=" * 70,
        ])
        
        with open(path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        return path
    
    def _generate_pdf(
        self,
        analysis: ContentAnalysis,
        path: str,
        include_full: bool,
        ai_content: Optional[GeneratedContent] = None
    ) -> str:
        """Generate PDF document - simplified robust version"""
        try:
            from fpdf import FPDF
        except ImportError:
            print("⚠️ PDF requires fpdf2: pip install fpdf2")
            return self._generate_text(analysis, path.replace('.pdf', '.txt'), include_full, ai_content)
        
        # Ultra-safe text cleaner
        def clean(text, maxlen=200):
            if not text:
                return ""
            s = str(text)[:maxlen]
            # Keep only basic ASCII printable chars
            s = ''.join(c if 32 <= ord(c) < 127 else ' ' for c in s)
            s = ' '.join(s.split())  # normalize whitespace
            return s
        
        try:
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            
            # Title
            pdf.set_font('Helvetica', 'B', 18)
            pdf.cell(0, 10, clean(analysis.title, 60), ln=True, align='C')
            pdf.set_font('Helvetica', '', 10)
            pdf.cell(0, 6, f"Words: {analysis.word_count} | Reading time: {analysis.reading_time_minutes} min", ln=True, align='C')
            pdf.ln(10)
            
            # Executive Summary
            pdf.set_font('Helvetica', 'B', 12)
            pdf.cell(0, 8, "EXECUTIVE SUMMARY", ln=True)
            pdf.set_font('Helvetica', '', 10)
            summary = clean(analysis.executive_summary, 600)
            for line in [summary[i:i+90] for i in range(0, len(summary), 90)]:
                pdf.cell(0, 5, line, ln=True)
            pdf.ln(5)
            
            # Simple Explanation (AI)
            if ai_content and ai_content.simplified_explanation:
                pdf.set_font('Helvetica', 'B', 12)
                pdf.cell(0, 8, "SIMPLE EXPLANATION", ln=True)
                pdf.set_font('Helvetica', '', 10)
                explanation = clean(ai_content.simplified_explanation, 400)
                for line in [explanation[i:i+90] for i in range(0, len(explanation), 90)]:
                    pdf.cell(0, 5, line, ln=True)
                pdf.ln(5)
            
            # Key Takeaways (AI)
            if ai_content and ai_content.key_takeaways:
                pdf.set_font('Helvetica', 'B', 12)
                pdf.cell(0, 8, "KEY TAKEAWAYS", ln=True)
                pdf.set_font('Helvetica', '', 10)
                for i, t in enumerate(ai_content.key_takeaways[:5], 1):
                    pdf.cell(0, 5, f"{i}. {clean(t, 80)}", ln=True)
                pdf.ln(5)
            
            # Key Sentences
            pdf.set_font('Helvetica', 'B', 12)
            pdf.cell(0, 8, "KEY SENTENCES", ln=True)
            pdf.set_font('Helvetica', '', 10)
            for i, s in enumerate(analysis.key_sentences[:5], 1):
                pdf.cell(0, 5, f"{i}. {clean(s, 85)}", ln=True)
            pdf.ln(5)
            
            # Key Concepts
            pdf.set_font('Helvetica', 'B', 12)
            pdf.cell(0, 8, "KEY CONCEPTS", ln=True)
            for c in analysis.concepts[:6]:
                pdf.set_font('Helvetica', 'B', 10)
                pdf.cell(0, 5, f"- {clean(c.term, 30)}", ln=True)
                pdf.set_font('Helvetica', '', 9)
                pdf.cell(0, 4, f"  {clean(c.definition, 80)}", ln=True)
            pdf.ln(5)
            
            # Study Questions
            pdf.add_page()
            pdf.set_font('Helvetica', 'B', 12)
            pdf.cell(0, 8, "STUDY QUESTIONS", ln=True)
            for i, q in enumerate(analysis.questions[:6], 1):
                pdf.set_font('Helvetica', 'B', 10)
                diff = (q.difficulty or "medium").upper()
                pdf.cell(0, 5, f"{i}. [{diff}] {clean(q.question, 70)}", ln=True)
                if q.answer_hint:
                    pdf.set_font('Helvetica', 'I', 9)
                    pdf.cell(0, 4, f"   Hint: {clean(q.answer_hint, 60)}", ln=True)
            pdf.ln(5)
            
            # FAQ (AI)
            if ai_content and ai_content.faq:
                pdf.set_font('Helvetica', 'B', 12)
                pdf.cell(0, 8, "FAQ", ln=True)
                for i, qa in enumerate(ai_content.faq[:4], 1):
                    q = qa.get('q', '') if isinstance(qa, dict) else str(qa)
                    a = qa.get('a', '') if isinstance(qa, dict) else ''
                    pdf.set_font('Helvetica', 'B', 10)
                    pdf.cell(0, 5, f"Q{i}: {clean(q, 70)}", ln=True)
                    pdf.set_font('Helvetica', '', 9)
                    pdf.cell(0, 4, f"A: {clean(a, 80)}", ln=True)
                    pdf.ln(2)
                pdf.ln(3)
            
            # Vocabulary (AI)
            if ai_content and ai_content.vocabulary:
                pdf.set_font('Helvetica', 'B', 12)
                pdf.cell(0, 8, "VOCABULARY", ln=True)
                for v in ai_content.vocabulary[:6]:
                    term = v.get('term', '') if isinstance(v, dict) else str(v)
                    meaning = v.get('meaning', '') if isinstance(v, dict) else ''
                    pdf.set_font('Helvetica', 'B', 10)
                    pdf.cell(40, 5, clean(term, 25))
                    pdf.set_font('Helvetica', '', 9)
                    pdf.cell(0, 5, clean(meaning, 60), ln=True)
                pdf.ln(3)
            
            # Related Topics
            pdf.set_font('Helvetica', 'B', 12)
            pdf.cell(0, 8, "RELATED TOPICS", ln=True)
            pdf.set_font('Helvetica', '', 10)
            topics = ", ".join(clean(t, 20) for t in analysis.related_topics[:6])
            pdf.cell(0, 5, topics, ln=True)
            
            # Footer
            pdf.ln(10)
            pdf.set_font('Helvetica', 'I', 8)
            if ai_content:
                pdf.cell(0, 4, "* Includes AI-generated content", ln=True, align='C')
            pdf.cell(0, 4, "Generated by EchoNotes", ln=True, align='C')
            
            pdf.output(path)
            return path
            
        except Exception as e:
            print(f"[PDF] Generation failed: {e}, falling back to text")
            # Fallback to text file
            return self._generate_text(analysis, path.replace('.pdf', '.txt'), include_full, ai_content)
    
    def _generate_docx(
        self,
        analysis: ContentAnalysis,
        path: str,
        include_full: bool,
        ai_content: Optional[GeneratedContent] = None
    ) -> str:
        """Generate Word document"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("⚠️ DOCX requires python-docx: pip install python-docx")
            return self._generate_text(analysis, path.replace('.docx', '.txt'), include_full, ai_content)
        
        doc = Document()
        
        # Title
        title = doc.add_heading(analysis.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Words: {analysis.word_count} | Reading: {analysis.reading_time_minutes} min").italic = True
        
        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(analysis.executive_summary)
        
        # AI Enhanced: Simple Explanation
        if ai_content and ai_content.simplified_explanation:
            doc.add_heading("Simple Explanation", level=1)
            doc.add_paragraph(ai_content.simplified_explanation)
        
        # Key Takeaways
        if ai_content and ai_content.key_takeaways:
            doc.add_heading("Key Takeaways", level=1)
            for i, takeaway in enumerate(ai_content.key_takeaways, 1):
                doc.add_paragraph(f"{i}. {takeaway}")
        
        # Key Sentences
        doc.add_heading("Key Sentences", level=1)
        for i, sent in enumerate(analysis.key_sentences[:5], 1):
            doc.add_paragraph(f"{i}. {sent}")
        
        # Key Concepts
        doc.add_heading("Key Concepts", level=1)
        for concept in analysis.concepts[:8]:
            p = doc.add_paragraph()
            p.add_run(f"{concept.term}: ").bold = True
            p.add_run(concept.definition[:200] if concept.definition else "")
        
        # AI Enhanced: Examples
        if ai_content and ai_content.examples:
            doc.add_heading("Real-World Examples", level=1)
            for i, example in enumerate(ai_content.examples, 1):
                doc.add_paragraph(f"Example {i}: {example}")
        
        # Study Questions
        doc.add_heading("Study Questions", level=1)
        for i, q in enumerate(analysis.questions[:8], 1):
            p = doc.add_paragraph()
            difficulty = q.difficulty.upper() if q.difficulty else "MEDIUM"
            p.add_run(f"{i}. [{difficulty}] ").bold = True
            p.add_run(q.question)
            if q.answer_hint:
                hint = doc.add_paragraph()
                hint.add_run(f"Hint: {q.answer_hint[:150]}").italic = True
        
        # AI Enhanced: FAQ
        if ai_content and ai_content.faq:
            doc.add_heading("Frequently Asked Questions", level=1)
            for i, qa in enumerate(ai_content.faq, 1):
                p = doc.add_paragraph()
                p.add_run(f"Q{i}: {qa.get('q', '')}").bold = True
                doc.add_paragraph(f"A: {qa.get('a', '')}")
        
        # AI Enhanced: Vocabulary
        if ai_content and ai_content.vocabulary:
            doc.add_heading("Vocabulary", level=1)
            for vocab in ai_content.vocabulary[:8]:
                p = doc.add_paragraph()
                p.add_run(f"{vocab.get('term', '')}: ").bold = True
                p.add_run(vocab.get('meaning', '')[:100])
        
        # Related Topics
        doc.add_heading("Related Topics", level=1)
        doc.add_paragraph(" | ".join(analysis.related_topics[:8]))
        
        # Action Items
        if analysis.action_items:
            doc.add_heading("Action Items", level=1)
            for item in analysis.action_items:
                doc.add_paragraph(f"☐ {item}")
        
        # AI indicator
        if ai_content:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("This document includes AI-generated content for enhanced understanding.").italic = True
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Generated by EchoNotes - Speech to Document System").italic = True
        
        doc.save(path)
        return path
